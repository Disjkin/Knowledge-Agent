"""文档加载器 - 支持PDF、Word、TXT、Markdown格式"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据模型"""
    content: str           # 文本内容
    metadata: dict = field(default_factory=dict)
    # metadata 包含：
    # - source: 文件路径
    # - file_type: 文件类型
    # - page: 页码（PDF）
    # - chunk_index: 分块索引（分割后）


class DocumentLoader:
    """统一文档加载接口"""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt", ".md"}

    def load(self, file_path: str) -> List[Document]:
        """根据文件扩展名选择合适的加载器"""
        path = Path(file_path)
        if not path.exists():
            logger.error(f"文件不存在: {file_path}")
            return []

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(path)
        elif suffix == ".docx":
            return self._load_docx(path)
        elif suffix == ".txt":
            return self._load_txt(path)
        elif suffix == ".md":
            return self._load_markdown(path)
        else:
            logger.warning(f"不支持的文件格式: {suffix}")
            return []

    def load_directory(self, dir_path: str) -> List[Document]:
        """加载目录下所有支持的文档"""
        path = Path(dir_path)
        if not path.is_dir():
            logger.error(f"目录不存在: {dir_path}")
            return []

        documents = []
        for file_path in sorted(path.rglob("*")):
            if file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                docs = self.load(str(file_path))
                documents.extend(docs)
                logger.info(f"已加载: {file_path.name} ({len(docs)}个文档片段)")

        logger.info(f"目录加载完成，共 {len(documents)} 个文档片段")
        return documents

    def _load_pdf(self, path: Path) -> List[Document]:
        """加载PDF文件"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            documents = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(Document(
                        content=text.strip(),
                        metadata={
                            "source": str(path),
                            "file_type": "pdf",
                            "file_name": path.name,
                            "page": i + 1,
                        }
                    ))
            return documents
        except Exception as e:
            logger.error(f"PDF加载失败 {path.name}: {e}")
            return []

    def _load_docx(self, path: Path) -> List[Document]:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)
            if content:
                return [Document(
                    content=content,
                    metadata={
                        "source": str(path),
                        "file_type": "docx",
                        "file_name": path.name,
                    }
                )]
            return []
        except Exception as e:
            logger.error(f"Word加载失败 {path.name}: {e}")
            return []

    def _load_txt(self, path: Path) -> List[Document]:
        """加载TXT文件"""
        try:
            # 尝试多种编码
            for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    content = path.read_text(encoding=encoding)
                    if content.strip():
                        return [Document(
                            content=content.strip(),
                            metadata={
                                "source": str(path),
                                "file_type": "txt",
                                "file_name": path.name,
                            }
                        )]
                    return []
                except UnicodeDecodeError:
                    continue
            logger.error(f"无法检测编码: {path.name}")
            return []
        except Exception as e:
            logger.error(f"TXT加载失败 {path.name}: {e}")
            return []

    def _load_markdown(self, path: Path) -> List[Document]:
        """加载Markdown文件"""
        try:
            content = path.read_text(encoding="utf-8")
            if content.strip():
                return [Document(
                    content=content.strip(),
                    metadata={
                        "source": str(path),
                        "file_type": "md",
                        "file_name": path.name,
                    }
                )]
            return []
        except Exception as e:
            logger.error(f"Markdown加载失败 {path.name}: {e}")
            return []
